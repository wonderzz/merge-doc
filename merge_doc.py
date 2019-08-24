# -*- coding: utf-8 -*-：
#    Author : wonder_zz
#    Time   : 2019/8/20  20:41
import os
import time
from docx import Document
from docxtpl import DocxTemplate
from win32com import client as wc


def ReSaveDoc(path, filename):
    """
    switch doc 2 docx
    """
    time1 = time.time()
    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(path + "\\" + filename)
    filename = filename.replace("doc", "docx")
    doc.SaveAs(path + "\\" + filename, 12, False, "",
               True, "", False, False, False, False)
    doc.Close()
    word.Quit()
    time2 = time.time()
    print("success change file " + filename + " in " + str(time2 - time1))


def ReSaveAllDoc(path):
    """
    find all doc file then resave them
    """
    filelist = []
    dirs = os.listdir()
    for f in dirs:
        filelist.append(str(f))
    for file in filelist:
        if (file.find(".doc") != -1) and (file.find(".docx") == -1):
            ReSaveDoc(path, file)
            print(file)
        else:
            continue


def combine_word_documents(path, files):
    """
    merge docx
    """
    merged_document = Document()
    merged_document.save(path + "\\" + 'merge.docx')
    try:
        for index, file in enumerate(files):
            doc = Document(path + "\\" + 'merge.docx')
            doc.add_paragraph('{{temp_name}}')
            doc.add_page_break()
            doc.save(path + "\\" + 'merge.docx')
            doc = DocxTemplate(path + "\\" + '合并文件.docx')
            sub = doc.new_subdoc()
            sub.subdocx = Document(path + "\\" + file)
            doc.render({'temp_name': sub})
            doc.save(path + "\\" + 'merge.docx')
    except Exception as e:
        pass


def MergeDocx(path):
    """
    find all the docx file
    """
    files = []
    filelist = []
    dirs = os.listdir()
    for f in dirs:
        filelist.append(str(f))
    for file in filelist:
        if (file.find(".docx") != -1) and (file.find("~$") == -1):
            files.append(file)
    combine_word_documents(path, files)


if __name__ == '__main__':
    path = os.getcwd()
    # Test(path)
    ReSaveAllDoc(path)
    MergeDocx(path)
